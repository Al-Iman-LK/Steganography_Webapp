from PIL import Image
import numpy as np
from django.core.files.base import ContentFile
import io
import PyPDF2
from docx import Document
import wave
import cv2
from pydub import AudioSegment
import tempfile
import os
import magic  # Add this import

class SteganographyProcessor:
    SUPPORTED_FORMATS = {
        'image': ['.png', '.jpg', '.jpeg'],
        'document': ['.pdf', '.doc', '.docx'],  # Added .doc
        'audio': ['.mp3', '.wav'],
        'video': ['.mp4']
    }

    MIME_TYPES = {
        'document': [
            'application/pdf',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/x-pdf',
            'application/vnd.ms-word',
            'application/octet-stream'  # Some systems may report this for .doc/.docx
        ]
    }

    def __init__(self, stegano_file):
        self.stegano_file = stegano_file
        self.mime = magic.Magic(mime=True)

    def _validate_file_type(self):
        """Validate file type and extension"""
        file_ext = os.path.splitext(self.stegano_file.original_file.name)[1].lower()
        file_type = self.stegano_file.file_type
        
        if file_ext not in self.SUPPORTED_FORMATS[file_type]:
            raise ValueError(f"Unsupported file format {file_ext} for type {file_type}")
        
        # Special handling for DOCX files
        if file_ext == '.docx':
            # DOCX files are actually ZIP files containing XML, so bypass MIME check
            return True
            
        # For other document types
        if file_type == 'document':
            mime_type = self.mime.from_buffer(self.stegano_file.original_file.read(1024))
            self.stegano_file.original_file.seek(0)
            
            if mime_type not in self.MIME_TYPES[file_type] and 'application/zip' not in mime_type:
                raise ValueError(f"Invalid document type detected: {mime_type}")
        
        return True

    def embed_message(self, message):
        """Embed a message in the file using appropriate steganography technique"""
        try:
            self._validate_file_type()
            
            if self.stegano_file.file_type == 'image':
                return self._process_image(message)
            elif self.stegano_file.file_type == 'document':
                return self._process_document(message)
            elif self.stegano_file.file_type == 'audio':
                return self._process_audio(message)
            elif self.stegano_file.file_type == 'video':
                return self._process_video(message)
        except Exception as e:
            self.stegano_file.status = 'failed'
            self.stegano_file.save()
            raise ValueError(f"Processing failed: {str(e)}")

    def _process_image(self, message):
        """Process image files for steganography"""
        try:
            # Open the image
            img = Image.open(self.stegano_file.original_file)
            # Convert image to numpy array
            img_array = np.array(img)

            # Basic LSB steganography implementation
            binary_message = ''.join(format(ord(i), '08b') for i in message)
            binary_message += '00000000'  # End marker

            if len(binary_message) > img_array.size:
                raise ValueError("Message too large for this image")

            # Embed the message
            idx = 0
            for i in range(img_array.shape[0]):
                for j in range(img_array.shape[1]):
                    for k in range(3):  # RGB channels
                        if idx < len(binary_message):
                            img_array[i, j, k] = img_array[i, j, k] & ~1 | int(binary_message[idx])
                            idx += 1

            # Save the processed image
            processed_img = Image.fromarray(img_array)
            output = io.BytesIO()
            processed_img.save(output, format='PNG')
            
            # Save to model
            self.stegano_file.processed_file.save(
                f'processed_{self.stegano_file.original_file.name}',
                ContentFile(output.getvalue()),
                save=True
            )
            self.stegano_file.status = 'completed'
            self.stegano_file.hidden_message = message
            self.stegano_file.save()

        except Exception as e:
            self.stegano_file.status = 'failed'
            self.stegano_file.save()
            raise e

    def _detect_file_type(self, file_path):
        """Detect the actual file type"""
        # Add DOCX mime type detection
        file_ext = file_path.lower().split('.')[-1]
        if file_ext == 'docx':
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        mime = magic.from_file(file_path, mime=True)
        return mime

    def _process_document(self, message):
        """Process document file with steganography"""
        temp_path = None
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(self.stegano_file.original_file.read())
                temp_path = temp_file.name

            file_ext = self.stegano_file.original_file.name.split('.')[-1].lower()
            
            if file_ext == 'pdf':
                return self._process_pdf(temp_path, message)
            elif file_ext == 'docx':
                return self._process_docx(temp_path, message)
            else:
                raise ValueError(f"Unsupported document type: {file_ext}")

        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass

    def _process_pdf(self, temp_path, message):
        """Process PDF files"""
        reader = PyPDF2.PdfReader(temp_path)
        writer = PyPDF2.PdfWriter()
        
        # Copy all pages
        for page in reader.pages:
            writer.add_page(page)
        
        # Add metadata with the hidden message
        writer.add_metadata({
            '/SteganoMessage': message.encode('utf-8')
        })
        
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        
        self.stegano_file.processed_file.save(
            f'processed_{self.stegano_file.original_file.name}',
            ContentFile(output.getvalue()),
            save=True
        )
        self.stegano_file.status = 'completed'
        self.stegano_file.hidden_message = message
        self.stegano_file.save()

    def _process_docx(self, temp_path, message):
        """Process DOCX files"""
        try:
            # Load the document from temp_path
            doc = Document(temp_path)
            
            # Add hidden paragraph with message
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"SteganoMessage:{message}")
            font = run.font
            font.hidden = True
            
            # Save to new file
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Save to model
            self.stegano_file.processed_file.save(
                f'processed_{self.stegano_file.original_file.name}',
                ContentFile(output.getvalue()),
                save=True
            )
            
            self.stegano_file.hidden_message = message
            self.stegano_file.status = 'completed'
            self.stegano_file.save()

        except Exception as e:
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    def _process_audio(self, message):
        """Process audio files for steganography"""
        temp_path = None
        wav_path = None
        try:
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as temp_file:
                temp_file.write(self.stegano_file.original_file.read())
                temp_path = temp_file.name

            # Convert MP3 to WAV using pydub
            if self.stegano_file.original_file.name.endswith('.mp3'):
                audio = AudioSegment.from_mp3(temp_path)
                wav_path = temp_path.replace('.mp3', '.wav')
                audio.export(wav_path, format='wav')
                temp_path = wav_path

            # Process WAV file
            audio_file = wave.open(temp_path, 'rb')
            params = audio_file.getparams()
            frames = bytearray(audio_file.readframes(audio_file.getnframes()))
            audio_file.close()

            # Embed message
            binary_message = ''.join(format(ord(i), '08b') for i in message)
            binary_message += '00000000'  # End marker

            if len(binary_message) > len(frames):
                raise ValueError("Message too large for this audio file")

            # Modify frames with message
            for idx, bit in enumerate(binary_message):
                if idx < len(frames):
                    frames[idx] = frames[idx] & ~1 | int(bit)

            # Create output WAV file
            output_wav = wave.open(temp_path, 'wb')
            output_wav.setparams(params)
            output_wav.writeframes(frames)
            output_wav.close()

            # Save the processed file
            with open(temp_path, 'rb') as f:
                self.stegano_file.processed_file.save(
                    f'processed_{self.stegano_file.original_file.name}',
                    ContentFile(f.read()),
                    save=True
                )

            self.stegano_file.status = 'completed'
            self.stegano_file.hidden_message = message
            self.stegano_file.save()

        except Exception as e:
            self.stegano_file.status = 'failed'
            self.stegano_file.save()
            raise ValueError(f"Audio processing failed: {str(e)}")
        finally:
            # Clean up temporary files
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
            if wav_path and os.path.exists(wav_path):
                try:
                    os.unlink(wav_path)
                except Exception:
                    pass

    def _process_video(self, message):
        """Process video files for steganography - frame only approach"""
        temp_path = None
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as temp_file:
                temp_file.write(self.stegano_file.original_file.read())
                temp_path = temp_file.name

            # Read video
            cap = cv2.VideoCapture(temp_path)
            if not cap.isOpened():
                raise ValueError("Could not open video file")

            # Get video properties
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            fps = cap.get(cv2.CAP_PROP_FPS)

            # Read first frame
            ret, frame = cap.read()
            if not ret:
                raise ValueError("Could not read video frame")

            # Prepare message
            binary_message = ''.join(format(ord(i), '08b') for i in message)
            binary_message += '00000000'  # End marker

            if len(binary_message) > (width * height):
                raise ValueError(f"Message too large. Maximum size: {width * height // 8} characters")

            # Create output frame
            modified_frame = frame.copy()

            # Embed message in blue channel LSBs
            message_idx = 0
            for i in range(height):
                for j in range(width):
                    if message_idx < len(binary_message):
                        # Set LSB in blue channel
                        modified_frame[i, j, 0] = (modified_frame[i, j, 0] & 254) | int(binary_message[message_idx])
                        message_idx += 1

            # Save as PNG to preserve LSBs
            frame_path = temp_path.replace('.mp4', '_frame.png')
            cv2.imwrite(frame_path, modified_frame)

            # Save to model
            with open(frame_path, 'rb') as f:
                self.stegano_file.processed_file.save(
                    f'processed_{self.stegano_file.original_file.name}.png',  # Save as PNG
                    ContentFile(f.read()),
                    save=True
                )

            self.stegano_file.status = 'completed'
            self.stegano_file.hidden_message = message
            self.stegano_file.save()

        except Exception as e:
            self.stegano_file.status = 'failed'
            self.stegano_file.save()
            raise ValueError(f"Video processing failed: {str(e)}")
        finally:
            if 'cap' in locals():
                cap.release()
            for path in [temp_path, frame_path] if 'frame_path' in locals() else [temp_path]:
                if path and os.path.exists(path):
                    try:
                        os.unlink(path)
                    except Exception:
                        pass
